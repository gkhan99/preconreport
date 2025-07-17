import streamlit as st
from openai import OpenAI, RateLimitError
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import base64
import tiktoken
import os
import time

# === CONFIG ===
client = OpenAI(api_key=st.secrets["api"]["key"])

# === HELPERS ===

def encode_image(image_path: str) -> str:
    with open(image_path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode("utf-8")
    return f"data:image/jpeg;base64,{b64}"

def build_prompt() -> str:
    return (
        "As a civil engineer, I have some photos and would like to classify them into different categories before starting a project. "
        "Find if it contains any visible cracks, peeling paint, possible water damage, visual discoloration, honeycombing, spalling or any other possible damage. "
        "If nothing, then just mention a statement about the image. Sound it technical and to the point."
    )

def estimate_cost(prompt_text: str, response_text: str, model: str="gpt-4o") -> float:
    enc = tiktoken.encoding_for_model(model)
    in_tokens  = len(enc.encode(prompt_text))
    out_tokens = len(enc.encode(response_text))
    return round(in_tokens * 0.005/1000 + out_tokens * 0.015/1000, 6)

def generate_report(image_paths: list[str], output_path: str="output_report.docx") -> tuple[float, str]:
    doc = Document()
    total_cost = 0.0

    # === LOGO BASE64 ===
    LOGO_BASE64 = """
    iVBORw0KGgoAAAANSUhEUgAAAOoAAADYCAMAAADS+I/aAAABAlBMVEX///8kMl4AAADzVycAGFL39/nyRgDzVSQiMF0SJVcdLVvzVSUAHVQcLFvzSADzUyAAEU8AFFAAAEqTl6gYKVnzTxgAF1EAE08LIVVQWHjzThTw8PLExs//+vgAHlQAD0+ws8D82tL+8/BeZYEOABHn5+j+7Of9492PlKb5q5rGyNADAAn6uar7y8D0ZTz0XzL3lH34oY14fpTb3OE+SG2xsbP2gWT6wLL1elv7z8X1b0sXDhmbmp2KiYwhGyNOS1F0cnb4mIL4ppMxPWZoboehpbOtrK9hX2RCP0U1MTiVlJfT09X3i3LyOwD0akVwdYvExMVram5KR00AADtHT3EeGSPXp+InAAAc6klEQVR4nNVdCVuyTBc2UARBsBIjzH3XbFHLyjJLrcwWe776/3/lOzOgojAIpNV7P9f1VIjj3JwzZ5tFn2/DyJycbvoj/giKr7L48tud+BlcMrx09Nud+BkkaPH18Lc7sVEUX9r6b4l6aXoxkfX5sod39xnDjYlS0eLt2cPT08PLxEa7uB5kzxnxbv5nAvc588q8FuuMKqvMbPAW4UbmOOM7rF9mjQ3UX2W4jXmxegx/C9nX+QBtXxy/1tEvRYaXT0RJ4mma0VU6w8g8zUtimxFfjaL2wRNRZZ6Xxcsf7bcXHAV5uX16jKR5KAbVCrqWYGhaUt9eVJ4PqviurBikeVFlJBnYLwgwc3p4+KJKdJDJmhv/S0hk3iSeV2UGSepSpOUTfBmoyqDWRZGnGTx+T1U6GCz5iseIsXlgZuG69pT+LjKvQZqmefX1Hv4oibR0hS+LPC8iIb1ItFrXLtAifhoqzdMW7RyqtPRXffJUMqLK08GjNhYd6C0va1dB0uhnRaVVZJ9LDB3EDNHTeLNo7k6eSzVbubiaWfXfRv1cYuQ7zPakEoQBqV8Hy4OFCQOYZtDLiCqyS3VgeI5eKE1/mePkrl2/A2FPh3BdVCVJEo//gkmuS6IUBHsjaZK9kGaGBn4VsXzf4BrifKhTbU/1ExRYH84ziDIywSpT1/5sM0EwaTIY6983UydMUBZVGHyyJp57WR+PYHymv17p9IGijEZxRgQtRy8g7ovmJ/OqSnzwqKI/rRIDDV/VKyq//Eh+FCU8gI5k8b7oK4J50YRZBzp6RjMV4ky8s9eQWUIP4UjS75gjUX+D8T71tUcSL13AzxPZylL/FC5fRfSjJCGGpaupe8gyU7s79zbnkm5wRex0wOzAYFRPKsBUs1OLAF6aT/K1RbgP+2h1eunHUcr4VF4TCPSkfsRAKBQ8xi8dB3l9XBWnrE9kGkc/GZGX8RhN0CgiUiFU0J7BIhKiJkrUli51MGgWN/4AMufMEXRb1V3iCwOGg56GB8BLuqqfgKwTSHb4EsPgEDEDAa7mLhMoBGYq4GiNwjrUdRTCKqwxGXBXEr5yLi1GVe2f0uY2SkV58V77PAj4mMrMkUAUi8JAbEFpVX1FAr5s1y9xRw0dLF5CnI9M8/xaljnGVqyt6n4Y7Jo+7tEVI7srqb4RZkVTjEaD95wanwzDIztbnMVGbzIKmNAgva+0M3aPHyJFY/Z+IQfF49PKCTLnWPhvU6HDINdVWseRxJyvX7CJU+Z12SKAbZ2NHYhskNOEganrWPaCYcRz28ykfYFazAIVo685FGU+KKvoSWkMVV4f/xdTS32k8wODwHybmQnn8OinJrU9NQ0QCKk6GbCXyIqeo0t6v7OrQptzlZHf3kSJ5hcig+ILA6FRENQftw3GXJMuBJI8kwX29aloVd4yoPwmsuDFReQRDlWUdmljBEbRNKKDuICXL45lOqgyVw6bfAUDLEFmIDFLws8evhwdX1Sy+ifrQ/VcQpb7/PjyqOSbvbKJkAJnXGDvGRigfJDBkivi54xfTiCXIUmvl/VDx7Fb5kVkRFFkrux9pRTEhOr62ChevOpOGD5e16CT9ZqnIKjmfYIBRwjCoLXoFKz/1C6VQNqi6vYji5lLW5OFAGNDLkGMBZ+PP+uifapZCIgssaMuHavM+TqDYxSzBF8k8bR+j6oH2OpnpokZINGubKbwBREwL0oqCoPRnxmwUVlMrC7izKcCkQsvvq4z6YFwnZewCcyoU+PzZg5f1486g/QoKF7hJ3k/05xDNSj7EhfwJCTmbq0pD0Qt4Fzwx0GkoDkAlIg5qmgHCs3qQXq38f6e+9SRe39/3x0eVJuFwIo3l84hqAqa/HpFlc5LMphG5nzdkfELKndpTy84TWFkmaFtVSferO7mPlvbEb+QVLhoNBoOh0MA+AF/cIog+Ldj/c/cbrUZJ7eSKFp8CDgAHkWPb+uPi7NoZGrqikJ3/AH3F+TPiVeH7629mMCFQ+yWLdhQmBP8e/3B7oEN32WAUQSu0kbKMBVVr4Th4MhWacbpQV8QuGjInuMS4zCX3O4P0mNHvUEDSmLWWlLMHp5WdFq0pOdfEP6qxDfE0wMlpkTZFaIk8lX8wiBdWNkvyO2ki7Xa/IqI5h40v3WJVPjlso4M75317eNha1uJeiJpQFSJ9XcPbPt1rxoDytLl9y0TqhqpQVoO4lYvUEwkQmSkWtrdaiPp51zpLBksjN5GldivIkSq0wkR3yENia/8TV1uA9NKhZbgH9KVEnI4NARzL2bNaZbD/m+Lc5FtVFAaTeuOgZefurnsG8Q0kizpjtcr+CAufaHqAtZYVMkqZeom0x8ftlKct8Fpj+j21q6FVT5EEbGmswlIMnjx5U6S5AvzjfYoztOFIoNTN0RQS4Kzomxh9JqDbWVNemsGy6VyJptcvBCn1uIc3L2IKjvHweUUaSVTdb56ISOClmTfcDFFw/2bafynPyPRTQh0jvB2a7gcUrX1mhYO4nAuWZFlgq0koKRK/IwrjE35UJIklIEnLNsJDMPCxgQ6R0gRrPTYpxXsxKB8XIToSXY3m5VAVSF1qsPIuwTVN2R+LxizF4vvcspmBToHlyxbOdsjSTov0rKktqVZpcsxEFdG53qCSrXo98SVHFSXqAbKyY2YIgJYzl82S/YoCAQTFyqEwzTjOhq+QgNd0wWIe2nppFQELeaXGxqmuJ/jqZNNlZfH7IuMC0x34Gclh6WeZa66DrfRZLfIqPyyeUv/nOoawSWHi329ZHgcSrRFSfc+xTs3/tUg10M0PUEHVXrB8jY//b9BFMAK/cUY6kKlsX8oiXrSfiEybrIdg1xLVwzDSPfGVwPvkfDvEEUI7Q2MQzbxptKioRRXpMUg48brYNuk2+7s0jqqdPSnB+kSootafM4EJWNBvs3w4oJoVuDKwHUBhVzsl3R3Dlb4NMbG9SCD69GJyvk5mt3KiCYb6oVrWlhvTO8RYf+CYNtI74q0CtEOmmW4lyUn8XB2ar/weF3iGnjf/m2SUwify072WIIgQA6CfNqqg9mcyjEjvh2S5NrkfnmUGhH2Lybvl2iy4eQUAqcXtOxkFdMLMYjKqselKVfGGGuVU78+So1gtxsLUtImdxNXELSvDide1CAKF2hJX0Zz9WqwZIF94bfJLUNpGZT4UtTWlWRFyHVW5XMlJnhcytbRAiTNvSQMuWmT/RP2aBHhqCGVPZaQMUJrFmlmlbO5k1F1tyTzFs/kIPUDuZp7sKm5JS7KqqyiCdvVTH1XkpjwJeigCMKsnC5Ekrt/a5gasFeedTJ7GpRQHXxWXysRJ5SA6qXvQsZLyiTVaMMaf8bHmCEMFigA0+kMVpshFubv5ODRvYqLg3eqcYnj4M8ZJCO4z7lcaSk4XZ+RQKsallbfzmKoDEMHZV4u4jlqwwPZV8zNK4rCcdFoOMQiELuBXgyFotEoun/5RQUBzVmFbNvQGsGzWRaNoM60ZllsVitx+tAqTZnnlxYvnrzOcp4LCBqkF1/2nuEN5ftPc/NcOZ1OD8uNRm6/1e/3gXnMb0JMiUFO2+/v594bDXT/UjtcGgEaaeSgjX5UgVYsmvEr0Eb/E24rWzSCEe3PuJY0i5p4YYI0H3w7NK7jgjxgZrBA/mjOh5Fp+ciO6da2qfIRsABcXLhHWJAcKyy3Yd3KYhuBPSu5G7hqIqXRaiCIElVGnSltVuSD8ynS7DEEEGgu72g2HbJv9Rj3XEwOznvpX0l1NeKWVJe43oFIaVk9OQGDLM8m0u7UxfnDClq/YdgmY8l0K7IOqn4PbfjiBF8Qbc1uyRyryOOgKbXEEXMxFVpC5ZcLxZl2fb6yYGC2SFiqqycFLbBJqlvc/vSWy1c0iaYZoMzcs6LA0WaGv5G0bvcXFZjo4ZXc9J47VbwwrwOpyMGg6eIMuzFCs56oLkvVE9UAOZhJzhKdY6so6UQOLnjZy6v2/HkcpEitRsvDYbW5bB2J3QsEmtXhcHfZ2ey6beRgOCzbJB2xodUbE7rWnkj6Nhgdp6o4y/SaNkUkCAmE7VQq1c+V7We3q8NGH+6LCYo5rYeQwI8baaxqpDxAjWwL1iHEDCnz3HPinGFU7EwrMi0al8tJvDp1soEtB7kMG44KKcJ0L8J+hFs9TcdGOf+WzZDoxxw0ghtKmlpB3pMXUZALZilIzwOKQ4PryTktroRby+3PkPY7bGMrmiM2MiTYRidduUSzE6qW60hBWr6aDs/SfNmvb+g8xN8hSmTgPJ0nG6lPF5myMjNNicvKpb6JBy/R0yZlJFXbG3AJwcW0gDommiQzYkQv66KXRFcdd9ETCG6mo/7wFc1noH1IuOaAQuNzCC1oVT2/Pz1igvxsNnbLRSqeJK1EiZOclQW2iVStQ0EC2LCuYpcimi1tq2gbUkbUF9iiKXZegtiY52dR/rubKqhAsp9jx0PVhqqbRrbmgz4hIq2tq2hdZPZVL/jeo43QtBbl6xZqHHHTukCS6oFzg0Km6qYRhFhae9+LDKLU90/MYoXii4jXnM0nddyoL4oFNkl11zoMJ4JVtJgE7W5nZNOOZzBX95X67FrZXesbpup2RoHTi00VBinr8lzMIgrb7sqDRKppF4/MnO1PqbqUKjSlF4fvX1WVsZ+nyrmcKF6LVIkD3j3VUF9/a6lSsZ98PHBbCSVStSwEEZAkmXHXCgyPLW1LcI6+2zo+kWr5l6huCc6ypbTrou/fo6pYpnMmuHM0CESqDRczWkSqbhqZwZFY3QuVTHXfxVAghlw5L/NijsTadz8PRabqwpQTG/FE1Uk1p+ou4rTvpSuqZUIjnqg6Eaub3s16uVmqnpaDsX1CczM0XWVM016uY6yumeqWn7zHQYMb9zDv5V+kGiZXcDS4yvhnvfyDCry1lbKfgXATn88RHRCa+02prjJM3poNb5Sq1yWN9oap4El/ycNiLVR3PPUJsG1Tn/alXdY2dJAVeB1j1TNVoglBWLex+12qIXIp3qP9/bNU7Sa8PervVmijVD0+f0AyTaTqKX4AhPcJDf6yVIk2xOdreVxGuFmpeqfKbpGYBrxuKNnIWI1XDxCGZQ4vAkMHDmiLtKzB4dejiwvGIqSAqep1ed1aqEYb4Ner6fSw8T7gopywE0ntYESwBRk0EHbTJKCVXo33Bm4qtS0IQpILs8R030sRR6O6jrG6FU1FItt7qR34l9pp5XK58u5BtVptxp2VxIwINMfj5m7uc2vdwSbZgTmNlkLhqCJEUqn+PpJbs+mcHF5mUU1Xh8NGbrAvJJOC3+8HiQLCQIckg7DXJb8skeqKh8eGOCUZ207tDwbDqtP1HxjN9G6jMWhtpSIpP2irwHHoHBHWCLQMkyM8IO9rflNuqbIsWjcS8X+CkjZdLPeKo4NgWjs7qVgSLVtdefYEYSLI5TymV6psmFNiAjdID8duBmEz3Wj0uRToppvd3oTZkU1QXRirSJTJ7chgd+hiMALJYfk9lvIrUQ872pW0ZZNup2wN2CFRnUkVnc+S3BoMD9wvTvyf4kqQCyDEJZ59DUGqhWpDazEUTe7tAEvXJDV4j5Y0Z22Bd+/7aExUC+lGP4U2CoY4v9JquDlbaL1U3y2b/PS+/XaBarya6+8p0RCMzO1kbriqRrlRqiFrx+o12F+gWm1sgTTZrTAXSzo5SmjDVAkBv5tFX5ZUm8PPJKLJcn5h4M7K+sbD/U/CS9+iGrVs8ntSPWhs7SkQb0WVVL/s7BysKQLpBjqLihTEfYcqwQ1+g2q0kUpq4uTeh+5MUGH4mcJnUYU2IdWtlKVufYPqVhSdTxHjXIrTVy2HItOzqAgW5M9RDXOp/q5d2dUC43dQgnls8J+gGlIiLcK5OURUB4J/MQb6+1RZLtUaunQq1fdlnhujap3FeaAK45N1q7fNhiBYxbQbocqylk269qthZbvh0g7Fh60I4eCijVAlmHV3MTAo7n7aHU/fQW6HfIZaeBPOhlAJdlPwDgvRskvFLZSTgl2UTazFbSDcd7xgiY2mcml3PH3p/VWH/21EgQn5qtMqhJIauLS4gbI/uTJt2ohUCTPnjqiGkoLLyM/XHDg6zXEjUiVsoHBQMQzFWml3PH3p/rYzc7cRqqSl46vMEhB1WTSJ7zo/5XATCkyqA69Y4Rbe3nddUGglnde/NiFVogOz8zbhvZyHyombZR9EqXqfSiZV0eyqo6HYp6cSkZvpqU1QJZSBwS4RVgixfrdjdEb1l6VK3NBivWifVVjSs1krVeJY9TxBCBEEsb5lNVijO3arf1ZRddHLDVC1WVFpXmHIRgbfqVb/slRtVhkGloYF6K7LJG0jVD0XguxOr1hMWcP+b+iuRnUdFtizVG3XUw4NGswKn98uzq+HqlepEpcHIRTm21bDEWe7ctZHlWRDPEvVbzv6poOLFfbXMd/iapEscfOJR6rkBVoYugaHUq5E2iTFX2tZD+xljwhu0N7SBHBszfVdFVMaqY0uvPM6Vu1O5ph2zk9ehWiBtMJtdpGsR6rEOaBZx3cUV8rbbAnshpdTeqS6cg9rIO3wm3M0lFOIzF+UKjEr94YqqwXNG6bqySytMEoukYvoSR+RqpsqenStVNmI+0WYRBhOuiaupnZTRSdK1dveNbtIyR0COcOBCmvZgE3snJf1VKz/OwnZAg44o2oqf44qRwpqXOM9slCaIUrVTS/XSZXdW9NIbbKm89wId7rZzULcKeKBqsMDBFaisbNcbSOf9uFiN8AaD0tgQ2shGm9ZnKC7joNN1kg15nXl5mLn/RZebi1U/aTo3DVVYpLvCo2IVamYaIHdKDCRqtvtayxxd40LxD+tNy0Tt7O42blDpOp2+4+wBptUVQghGpGqmxPc1nQMHHlKyoTnPMklkb82grz918XWdaLaxV2d0zZX32t733rWyec/rF8akCfVFWKjLhIwhdQGaSrJGrMvf5lQE1+BGB6On2qd7gdllbDGW2Q7aFOvcn46JTG9d5ewzs/z/Hr23VBA1wqFR6pD3QR8o3/m15p263BsAk7nx4USrbgrb8NGpwpWoJ6BacdSRQvUF/WB9PyBMo2atClAMra+QzbugbCzr1hjOZsDkgqkBWzmZnZmViNA5akP31PPssVHXXPH1MPSK2WLKU42HOY4TvBvc6SvSNXQULb9SfRtBGELNQyhRpSksB1u2PnCeIPbFnAjIXtdTqXnb5o8XfualoMRixz/vKaul3prYZAi/c9cubx74KAQFWgepMvlxuenyYLvtFAjw3TVQR4y1hpp2U0sJ5dGUjz/SGjtBg/iSb62OJZzFh6DfHasHVZ/xcBqFMgntnIGy6b175n4DG+px5t/ta/uwkXLg/c9nbu/FqrkE9oN+xzGeWq0Qt8eqVqeelqQmHXTf5Bq2PB9Ch0KORP7hiY3N0ueyOobMv4i1VBo3qUxNUaes+tylMW3rDzqn6MaUua8mr0n+P+Mul3Z2Nmo1j0ztN234PrXqLKCQYLdTh71v7BSqGfU1+MNZeS6ZY7upmWqOKB5cICPcNg1I42uN+GegBVV/TsyAloj2jEQ1o1U541YLvoMR42+fdztUB9O6mijbqcHfA23WsS/EEAgKBFALAkQtF37S1DQln4B7hEgWAgvb05g0U74cAheT0X8AmqC0Ai8EINGkmGLRjDT8JIAb6hOx4HH7z7FqevxQpgRsLBN9t/DZHm79VXnrZBujpq/nmDSzTsYYh8dyH86OA4+mw3snPfN2RsH92nQwPE/ChuagJP5xAI4pBuqh5hSc9/07n3P/YaRNGZ/Y6pW+zIaGntMgCAK+K8pYwg5/JvficdGFiYFaj1f4LEzckrVN8FP5YbC9YiZIlT3fvH7kUlgjbkM0kPU3S6xZmSNG6p2g9880+FC/w99+aoGLjR1MpNHZIaua0gPnyxqCzbo5XFKNw6MDKHkwOXJ7ZuGPzcTX5dCpMdU53HyTEhSiTLFSevZUsycXr1d5ucQMq6XG4+bj7jfecpUWViBJ3Q/OJ2l3LZgMWPzS1D6C8ILaEZ0cvPsJZXWmC5pQ/lvfA8ru2eoOGByZ1R+RdpGxrjWQSHE7e1i3NHsO6uMbRRc1LBtYkIVbs5Q+a/27JHqI4VyoUDeNMrLsV8eseG9hSLS7Venhvzjc225KuYYPWTfJpjwIgqfXs/GWwuSrcUa5aTW1SLYXu1bi9Jv4EmNu72lqwfR3zJPLGea/X/+96UP0+/NNo6o+DNV+zJNCJRjng9B+g6ie42lLylF4/OMqj2ceVXeKQrUaER9WXmqeMN8CsCmETZtnY1TOLi/qeUd1Ffs0ezgeRx4cI+9ZckWBrEf/arzcCxnnki46VDIQ9x0PXuaGeI32Nc8Ut1HyjTX0xzsOJ1O+Taiqdzc6Jzh5GWCLOZoOcz5JiZUHs2AmBW58B77EbLc3sAg0QAOFwr4/3G+5jIStEcv/0E9jimr5xffFZIbDqBCyk55cYw+1GoF9D/y+aO8eerwG7imxvFbqgMP8bmnBU/G1tOfkc3ld2x0u2VeyDHCM6e3ndH4oXP9vS1eSwh8jQpg0gsoeJqcAc1naiEIaza4pOdzPO14hhS/5YkiYwoJtFDrUJ4jJBIKIwrbpxuwyFQPuD0tf0I6JyhrZhtSYvtpc0WhgC49dpBNat4+OS4kOcdDp4ktwbXvLO57sHqUgeG+IKyNbUjZ/rQ6lwEiXRTWF6jaBkga8NHRIuIvyjqJiB/kon7u21YqxPm5/fQCz8lHF8t3RF338sgVgH5tkmlTXzLykM+DEHvWKwOqu58RgQt5lW6IEyKt5S8UhvHTwcnkMwqOeqgmGKBG6zS8Jox7+Ec3/4ynus4erD8tXi3vRwQl6i7ZY0NRxR/ZL1fNatvN69p6i9S2Wav5TDWDzeAsj8qP/7qdW8qmvDxON3JKzBnhUJQT/ML+e3psOf8woaal3Q9ULhib47ZNoYuiFFDmJiJtf2uzOmzsR/0RvMoFLXOZzrqwbCgUjuLZpz0YmO/DA5szHsdUd3T7dPusFYEeRp21bUJYhbNRHA2YW7RGhIK/C9fP9toUiI/Tu+VyebDfaikpjFirtZ+DS+X0mLw8boZrcHFUHtUvH5CzW+MeGgeIU2ghzCMKWiZUjaKeJ7NljJNN9CSA+OIFKRs1RlYofOSRc0Mxy1fnIzCGhFFjGKDclZzJiN+M/hlNwYNlDP5DuEGFpzFViyP59rRrkxrlsbX4aF7GKiBj8O96NDV7gfEZ5JE/NkTNmNw+4NkD+HW2sOuauARsFZqzit0EOc5/X2gSRl9EBfqzanXOT+AJwuObztQUP0Lo6LsejZtnzrc2x3uaJgQCvsk/cGQPQAtChA+Kepzmomc/4kZXIQBd+vrq6X9h8fbyI10MYyeEIdVGM523EPAFkCe7RpHfV6c3/gUjZAccMXWni4zBPqHwAkbuBKfMj9T1U/5W52vqd2GCpfXV7VIUquPe4PSscIMs2w31OHm4pVyeIbhJQAbbu/43m6m+zqPKHYWKzyjG8D11uw+PHRTIIeFTj8D2A5WVn1GFb0R94ZTw4Sn/MHvvM/XYw9a8B/d//Am9nQHs5Dz0fkQ5D0RScfivBh2u5Se+QA2R/ld7HD9TXwEU8aABDmoAHjhOIdn3aigrnNTwgP/X+dIeXOAXLS4R86DhDNVQz1DMepbH+QcwwXTGiL5vhGpeI2oSp77Q3c9dqgvPApLQng+b2YL2o/cbJDxh8thDY+0DJyEBGL/UNIa8Rfb5gepNcJ0Eks/AI8pVzmojZL9q2ogv/Gzg931Meg8oaKTw/BEuGzR9gQ4mUxvdIhFD8IEKftc4zKR6IFxCUvifQHzygFRZMzfd55EW291ou3kCnXyX+kf10H1nk784ML1ifPPY0+K8+McH9iHx6+vmT6nr/wHZzG7XNLGBXQAAAABJRU5ErkJggg==
    """  # Replace with your actual base64 logo string

    logo_path = "temp_logo.png"
    try:
        with open(logo_path, "wb") as f:
            f.write(base64.b64decode(LOGO_BASE64.strip()))
    except Exception as e:
        st.warning(f"⚠️ Failed to decode or save logo: {e}")
        logo_path = None

    # === Add logo to header (repeats on every page) ===
    if logo_path and os.path.exists(logo_path):
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(2), height=Cm(2))

    # === Track images per page ===
    images_on_page = 0

    for idx, img_path in enumerate(image_paths, start=1):
        data_uri = encode_image(img_path)
        retries = 3
        delay = 5
        success = False

        for attempt in range(retries):
            try:
                resp = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": build_prompt()},
                                {"type": "image_url", "image_url": {"url": data_uri}},
                            ],
                        }
                    ],
                    max_tokens=200,
                )
                comment = resp.choices[0].message.content.strip().capitalize()
                if not comment.endswith("."):
                    comment += "."
                total_cost += estimate_cost(build_prompt(), comment)
                success = True
                break
            except RateLimitError:
                st.warning(f"⚠️ Rate limit hit. Retrying in {delay} sec... (Attempt {attempt+1}/{retries})")
                time.sleep(delay)
                delay *= 2
            except Exception as e:
                comment = f"⚠️ Error analyzing image: {e}"
                break

        if not success:
            comment = comment if 'comment' in locals() else "⚠️ Could not analyze image after retries."

        # === Add image & comment ===
        p_num = doc.add_paragraph(f"Image No.: {idx}")
        p_num.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p_img = doc.add_paragraph()
        run = p_img.add_run()
        run.add_picture(img_path, width=Cm(15), height=Cm(7.5))
        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        p_txt = doc.add_paragraph(f"Assessment: {comment}")
        p_txt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("")

        images_on_page += 1
        if images_on_page == 2:
            doc.add_page_break()
            images_on_page = 0

    doc.save(output_path)
    return total_cost, output_path

# === STREAMLIT PAGES ===

def login_page():
    st.title("🔐 Hirani Pre-Con Report Login")
    user = st.text_input("Username")
    pwd = st.text_input("Password", type="password")

    if st.button("Log in"):
        if (user == st.secrets["auth"]["username"]
         and pwd == st.secrets["auth"]["password"]):
            st.session_state.logged_in = True
            st.success("Login successful! Please upload images.")
        else:
            st.error("Invalid username or password")

def image_analysis_page():
    st.title("🖼️ Image Analysis & Report Generation")

    uploaded = st.file_uploader(
        "Upload images (.jpg/.jpeg/.png)",
        accept_multiple_files=True
    )
    if not uploaded:
        st.info("Please upload one or more images.")
        return

    os.makedirs("temp_images", exist_ok=True)
    image_paths: list[str] = []
    for f in uploaded:
        ext = os.path.splitext(f.name)[1].lower()
        if ext not in (".jpg", ".jpeg", ".png"):
            st.error(f"Invalid extension on `{f.name}` – only .jpg/.jpeg/.png allowed.")
            continue
        path = os.path.join("temp_images", f.name)
        with open(path, "wb") as out:
            out.write(f.getbuffer())
        image_paths.append(path)

    if not image_paths:
        st.warning("No valid images to process.")
        return

    st.write("✅ Saved images:", image_paths)
    if st.button("Generate Report"):
        cost, docx_file = generate_report(image_paths)
        st.success("Report generated!")
        st.write(f"• File: `{docx_file}`")
        st.write(f"• Estimated API cost: ${cost}")
        with open(docx_file, "rb") as docf:
            st.download_button("📥 Download .DOCX", docf, file_name=os.path.basename(docx_file))

# === ENTRY POINT ===

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if st.session_state.logged_in:
    image_analysis_page()
else:
    login_page()
